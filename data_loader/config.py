"""加载器配置——修改参数只需改这里。"""

import warnings
from pathlib import Path

from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document

# langchain-community 正在逐步迁移到独立包，但目前仍是这些 Loader 的主要来源
with warnings.catch_warnings():
    warnings.filterwarnings(
        "ignore", message=".*langchain-community.*is being sunset.*"
    )
    from langchain_community.document_loaders import (
        PyPDFLoader,
        TextLoader,
    )


class SimpleMarkdownLoader(TextLoader):
    """轻量级 Markdown 加载器——使用 TextLoader 读取原始内容。

    UnstructuredMarkdownLoader 依赖太重（需要 spaCy 模型），
    对于大多数场景，直接以纯文本读取 Markdown 并标注格式即可。

    编码由 LOADER_ARGS[".md"] 控制，默认 utf-8。
    """

    pass


class SimpleDocxLoader(BaseLoader):
    """轻量级 DOCX 加载器——使用项目已声明的 python-docx 依赖。"""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> list[Document]:
        from docx import Document as DocxDocument

        docx = DocxDocument(self.file_path)
        paragraphs = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
        tables = []
        for table in docx.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables.append(" | ".join(cells))

        content = "\n".join(paragraphs + tables)
        return [Document(page_content=content, metadata={"source": self.file_path})]


# 文档存储目录，默认项目根目录下的 Data/docs/ 文件夹
DOCUMENT_DIR = Path(__file__).resolve().parent.parent / "Data" / "docs"

# ---------------------------------------------------------------------------
# 文件格式 → 加载器 映射表
# ---------------------------------------------------------------------------
# 键：文件扩展名（小写），值：对应的 Loader 类
# 新增格式只需在此添加一条记录
LOADER_MAPPING = {
    ".txt": TextLoader,
    ".md": SimpleMarkdownLoader,
    ".pdf": PyPDFLoader,
    ".docx": SimpleDocxLoader,
}

# 每种加载器的额外参数（无参数则为空字典）
LOADER_ARGS = {
    ".txt": {"encoding": "utf-8"},
    ".md": {"encoding": "utf-8"},
}
